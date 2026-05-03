"""
CV Parser Service - Extract and analyze CV/Resume content
"""
import re
from typing import Dict, Any, List, Optional
from pathlib import Path

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from docx import Document
except ImportError:
    Document = None


class CVParserService:
    # Common technical skills
    TECH_SKILLS = {
        "languages": ["python", "javascript", "typescript", "java", "c++", "c#", "go", "rust", "ruby", "php", "swift", "kotlin", "scala", "r"],
        "frontend": ["react", "angular", "vue", "svelte", "next.js", "nuxt", "html", "css", "sass", "tailwind", "bootstrap"],
        "backend": ["node.js", "express", "fastapi", "django", "flask", "spring", "laravel", "rails", "asp.net"],
        "databases": ["mysql", "postgresql", "mongodb", "redis", "elasticsearch", "sqlite", "oracle", "sql server", "mariadb"],
        "cloud": ["aws", "azure", "gcp", "docker", "kubernetes", "terraform", "jenkins", "circleci", "github actions"],
        "tools": ["git", "jira", "confluence", "figma", "postman", "webpack", "vite", "npm", "yarn"]
    }
    
    # Email and phone patterns
    EMAIL_PATTERN = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    PHONE_PATTERN = r'[\+]?[(]?[0-9]{1,3}[)]?[-\s\.]?[(]?[0-9]{1,4}[)]?[-\s\.]?[0-9]{1,4}[-\s\.]?[0-9]{1,9}'
    
    # Section headers
    SECTION_HEADERS = {
        "experience": ["experience", "work experience", "employment", "work history", "professional experience"],
        "education": ["education", "academic", "qualifications", "academic background"],
        "skills": ["skills", "technical skills", "competencies", "technologies", "expertise"],
        "projects": ["projects", "personal projects", "portfolio", "notable projects"],
        "certifications": ["certifications", "certificates", "licenses", "credentials"],
        "summary": ["summary", "profile", "objective", "about me", "professional summary"]
    }
    
    async def parse_file(self, file_path: str) -> Dict[str, Any]:
        """Parse CV file and extract text content."""
        path = Path(file_path)
        
        if not path.exists():
            return {"error": "File not found", "text": ""}
        
        ext = path.suffix.lower()
        
        if ext == ".pdf":
            return await self._parse_pdf(file_path)
        elif ext in [".docx", ".doc"]:
            return await self._parse_docx(file_path)
        elif ext == ".txt":
            return await self._parse_txt(file_path)
        else:
            return {"error": f"Unsupported file type: {ext}", "text": ""}
    
    async def _parse_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extract text from PDF file."""
        if not PyPDF2:
            return {"error": "PyPDF2 not installed", "text": ""}
        
        try:
            text = ""
            with open(file_path, "rb") as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    text += page.extract_text() + "\n"
            return {"text": text.strip(), "pages": len(reader.pages)}
        except Exception as e:
            return {"error": str(e), "text": ""}
    
    async def _parse_docx(self, file_path: str) -> Dict[str, Any]:
        """Extract text from DOCX file."""
        if not Document:
            return {"error": "python-docx not installed", "text": ""}
        
        try:
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            return {"text": text.strip(), "paragraphs": len(doc.paragraphs)}
        except Exception as e:
            return {"error": str(e), "text": ""}
    
    async def _parse_txt(self, file_path: str) -> Dict[str, Any]:
        """Read text file."""
        try:
            with open(file_path, "r", encoding="utf-8") as file:
                text = file.read()
            return {"text": text.strip()}
        except Exception as e:
            return {"error": str(e), "text": ""}
    
    def extract_contact_info(self, text: str) -> Dict[str, Any]:
        """Extract contact information from CV text."""
        emails = re.findall(self.EMAIL_PATTERN, text)
        phones = re.findall(self.PHONE_PATTERN, text)
        
        return {
            "email": emails[0] if emails else None,
            "phone": phones[0] if phones else None,
            "all_emails": emails,
            "all_phones": phones
        }
    
    def extract_skills(self, text: str) -> Dict[str, List[str]]:
        """Extract technical skills from CV text."""
        text_lower = text.lower()
        found_skills = {
            "languages": [],
            "frontend": [],
            "backend": [],
            "databases": [],
            "cloud": [],
            "tools": [],
            "all": []
        }
        
        for category, skills in self.TECH_SKILLS.items():
            for skill in skills:
                # Use word boundaries for more accurate matching
                pattern = r'\b' + re.escape(skill) + r'\b'
                if re.search(pattern, text_lower):
                    found_skills[category].append(skill)
                    found_skills["all"].append(skill)
        
        return found_skills
    
    def extract_sections(self, text: str) -> Dict[str, str]:
        """Extract different sections from CV."""
        sections = {}
        lines = text.split("\n")
        current_section = "header"
        section_content = []
        
        for line in lines:
            line_lower = line.lower().strip()
            found_section = None
            
            for section, headers in self.SECTION_HEADERS.items():
                if any(header in line_lower for header in headers):
                    found_section = section
                    break
            
            if found_section:
                # Save previous section
                if section_content:
                    sections[current_section] = "\n".join(section_content)
                current_section = found_section
                section_content = []
            else:
                section_content.append(line)
        
        # Save last section
        if section_content:
            sections[current_section] = "\n".join(section_content)
        
        return sections
    
    def extract_experience(self, text: str) -> List[Dict[str, Any]]:
        """Extract work experience entries."""
        sections = self.extract_sections(text)
        experience_text = sections.get("experience", "")
        
        # Simple extraction - in production, use NLP
        experiences = []
        
        # Look for date patterns (e.g., "2020 - 2023", "Jan 2020 - Present")
        date_pattern = r'(\d{4}\s*[-–]\s*(?:\d{4}|[Pp]resent|[Cc]urrent))'
        
        lines = experience_text.split("\n")
        current_exp = {}
        
        for line in lines:
            line = line.strip()
            if not line:
                if current_exp:
                    experiences.append(current_exp)
                    current_exp = {}
                continue
            
            # Check for dates
            dates = re.findall(date_pattern, line)
            if dates:
                current_exp["period"] = dates[0]
                current_exp["title"] = line.replace(dates[0], "").strip(" -–")
            elif current_exp and not current_exp.get("description"):
                current_exp["description"] = line
        
        if current_exp:
            experiences.append(current_exp)
        
        return experiences
    
    def extract_education(self, text: str) -> List[Dict[str, Any]]:
        """Extract education entries."""
        sections = self.extract_sections(text)
        education_text = sections.get("education", "")
        
        education = []
        
        # Look for degree keywords
        degree_keywords = ["bachelor", "master", "phd", "diploma", "certificate", "bsc", "msc", "ba", "ma", "mba"]
        
        lines = education_text.split("\n")
        
        for line in lines:
            line_lower = line.lower().strip()
            for keyword in degree_keywords:
                if keyword in line_lower:
                    education.append({
                        "degree": line.strip(),
                        "raw": line
                    })
                    break
        
        return education
    
    async def analyze_cv(self, file_path: str) -> Dict[str, Any]:
        """Complete CV analysis."""
        # Parse file
        parsed = await self.parse_file(file_path)
        
        if parsed.get("error"):
            return {
                "error": parsed["error"],
                "success": False
            }
        
        text = parsed.get("text", "")
        
        if not text:
            return {
                "error": "Could not extract text from CV",
                "success": False
            }
        
        # Extract all components
        contact = self.extract_contact_info(text)
        skills = self.extract_skills(text)
        sections = self.extract_sections(text)
        experience = self.extract_experience(text)
        education = self.extract_education(text)
        
        # Calculate basic scores
        skill_count = len(skills.get("all", []))
        exp_count = len(experience)
        edu_count = len(education)
        
        format_score = min(100, 50 + (10 if contact.get("email") else 0) + (10 if contact.get("phone") else 0) + (len(sections) * 5))
        content_score = min(100, 30 + (skill_count * 3) + (exp_count * 10) + (edu_count * 10))
        
        return {
            "success": True,
            "text": text,
            "contact_info": contact,
            "skills": skills,
            "sections": list(sections.keys()),
            "experience": experience,
            "education": education,
            "metrics": {
                "word_count": len(text.split()),
                "skill_count": skill_count,
                "experience_count": exp_count,
                "education_count": edu_count
            },
            "preliminary_scores": {
                "format_score": format_score,
                "content_score": content_score
            }
        }


# Singleton instance
cv_parser = CVParserService()
